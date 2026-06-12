"""Generate (a) a Word .docx with headings and (b) a PDF with NO table of contents,
to exercise Word parsing and the AI semantic-segmentation fallback."""
import fitz
from docx import Document

# (a) Word document with Heading styles -> outline path for docx
doc = Document()
doc.add_heading("Newton's Laws of Motion", level=1)
doc.add_paragraph("An object at rest stays at rest unless acted on by an external force.")
doc.add_paragraph("This is the principle of inertia, the first of Newton's three laws.")
doc.add_heading("Thermodynamics", level=1)
doc.add_paragraph("The first law of thermodynamics states that energy cannot be created or destroyed.")
doc.add_paragraph("The second law states that entropy of an isolated system always increases.")
doc.add_heading("Electromagnetism", level=1)
doc.add_paragraph("Electric charges produce electric fields that exert forces on other charges.")
doc.add_paragraph("A changing magnetic field induces an electric current, known as induction.")
doc.save("test_physics.docx")
print("wrote test_physics.docx")

# (b) PDF with NO TOC -> forces AI semantic segmentation + LLM labeling
pdf = fitz.open()
blocks = [
    "The French Revolution began in 1789 and radically transformed France's political system.",
    "The storming of the Bastille became a symbol of the uprising against the monarchy.",
    "The Industrial Revolution introduced mechanized manufacturing in the 18th and 19th centuries.",
    "Steam power and factories dramatically increased production and reshaped society.",
    "World War I lasted from 1914 to 1918 and involved many of the world's great powers.",
    "Trench warfare characterized much of the conflict on the Western Front.",
]
for b in blocks:
    page = pdf.new_page()
    page.insert_text((72, 100), b, fontsize=12)
pdf.save("test_history_notoc.pdf")  # no set_toc -> no bookmarks
pdf.close()
print("wrote test_history_notoc.pdf")
